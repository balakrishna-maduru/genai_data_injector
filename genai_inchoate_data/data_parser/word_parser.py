import os
from PIL import Image
from io import BytesIO
from docx import Document
from genai_inchoate_data.data_parser.data_parser import DataParser


class WordParser(DataParser):
    def __init__(self, docx_path):
        self.docx_path = docx_path

    def extract_metadata(self):
        metadata = {}
        metadata['file_name'] = os.path.basename(self.docx_path)
        metadata['file_location'] = os.path.dirname(os.path.abspath(self.docx_path))
        # Add more metadata extraction based on your Word file structure
        return metadata

    def extract_text(self):
        document = Document(self.docx_path)
        text_content = []
        for paragraph in document.paragraphs:
            text_content.append(paragraph.text)
        return '\n'.join(text_content)

    def extract_images(self, output_folder):
        document = Document(self.docx_path)
        image_counter = 1
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                image_stream = rel.target_part.blob
                image = Image.open(BytesIO(image_stream))
                image_path = os.path.join(output_folder, f"image_{image_counter}.png")
                image.save(image_path)
                image_counter += 1

    def extract_tables(self):
        document = Document(self.docx_path)
        tables = []

        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return tables
